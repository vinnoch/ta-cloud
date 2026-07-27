from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from build_launch_report import (
    GRAY,
    GREEN,
    INK,
    LIGHT,
    NAVY,
    PALE,
    TEAL,
    WHITE,
    add_field,
    add_heading,
    add_para,
    bullet,
    callout,
    document_settings,
    page_break,
    set_cell_text,
    set_width,
    shade,
    table,
)


OUT = Path(__file__).resolve().parent
DOCX = OUT / "TACLOUD_Launch_Readiness_and_Project_Handover_Report.docx"
SEQUENCE_PNG = OUT / "TACLOUD_Thesis_Lifecycle_Sequence.png"
LOGIN_SCREENSHOT = OUT / "TACLOUD_Login_Page.png"
COVER_VISUAL = OUT / "TACLOUD_Cover_Login_Visual.png"
UKWK_LOGO = Path(__file__).resolve().parents[3] / "public/images/ukwk-logo.png"


TEST_SCENARIOS = [
    ("1", "Test Environment and Data", "Confirmed that the test database, migrations, accounts, roles, and sample data were ready.", "Passed"),
    ("2", "Application Baseline", "Checked the Laravel test suite, application routes, Blade pages, and frontend build.", "Passed"),
    ("3", "Browser Test Setup", "Confirmed that TestSprite could open TACLOUD and run the accepted browser scenarios.", "Passed"),
    ("4", "Login and Session Security", "Tested valid and invalid login, logout, rate limiting, session handling, and institutional Google accounts.", "Passed"),
    ("5", "Role Access", "Confirmed that Kaprodi, Dosen, and Mahasiswa can only access functions allowed for their roles.", "Passed"),
    ("6", "Record Ownership", "Checked that users cannot view or change thesis, guidance, or document records belonging to other users.", "Passed"),
    ("7", "CSRF Protection", "Confirmed that unauthorized form submissions without a valid security token are rejected.", "Passed"),
    ("8", "Input Validation", "Tested invalid and manipulated form data to ensure restricted fields cannot be changed.", "Passed"),
    ("9", "File Security", "Checked file type, upload path, ownership, thesis phase, preview, and download permissions.", "Passed"),
    ("10", "XSS Protection", "Tested whether harmful text or scripts could be stored or displayed through application inputs.", "Passed"),
    ("11", "Search and Sorting Safety", "Confirmed that search, filtering, and table sorting only use accepted values.", "Passed"),
    ("12", "Browser Security Headers", "Checked response headers used to reduce browser-based security risks.", "Passed"),
    ("13", "Mahasiswa Workflow", "Tested proposal management, document uploads, grade access, and final-document submission.", "Passed"),
    ("14", "Dosen Workflow", "Tested guidance records, grading, grade locking, and sidang request submission or cancellation.", "Passed"),
    ("15", "Kaprodi Workflow", "Tested proposal review, assignments, sidang approval, grade unlock, and thesis completion.", "Passed"),
    ("16", "Complete Thesis Lifecycle", "Confirmed that a thesis can move through the main phases and appear in the public library after completion.", "Passed"),
    ("17", "Notifications and Privacy", "Checked notification recipients, notification links, read status, and private data access.", "Passed"),
    ("18", "Accessibility", "Tested keyboard navigation, focus handling, modal dialogs, labels, and Escape-key behavior.", "Passed"),
    ("19", "Production Routes and Cache", "Confirmed that routes are unique and the application can build route and configuration caches.", "Passed"),
    ("20", "Dependencies and Release Checks", "Checked Composer and npm security reports, frontend build, optimization, migrations, and the final test suite.", "Passed"),
]


QA = {
    "Kaprodi": [
        ("1.1", "Dashboard", "Dashboard widgets, counts, and statistics loaded correctly.", "Passed", "Verified"),
        ("1.2", "Program Study Master Data", "Create, view, edit, and delete Program Study records.", "Passed", "Verified"),
        ("1.3", "Dosen Master Data", "Dosen record management and CSV import.", "Passed", "Verified"),
        ("1.4", "Mahasiswa Master Data", "Mahasiswa record management and CSV import.", "Passed", "Verified"),
        ("1.5", "Academic Year", "Create, view, edit, and delete academic-year records.", "Passed", "Verified"),
        ("1.6", "Academic Period", "Create, view, edit, and delete academic periods.", "Passed", "Verified"),
        ("1.7", "Grading Format", "Create, view, edit, and duplicate grading formats.", "Passed", "Verified"),
        ("1.8", "Thesis Management", "View the thesis list and open thesis details.", "Passed", "Verified"),
        ("1.9", "Reviewer Assignment", "Assign supervisors and examiners to a thesis.", "Passed", "Verified"),
        ("1.10", "Proposal Review", "Review, approve, or reject a student proposal.", "Passed", "Verified"),
        ("1.11", "Sidang Request Review", "Review, approve, or reject a sidang request.", "Passed with Remark", "The tester saw an error after approval, although the status changed. The corrected flow is technically verified; an identical Citra retest is not documented."),
        ("1.12", "Grade Recap and Unlock", "View final grade recap and process grade-unlock requests.", "Passed", "Verified"),
        ("1.13", "Document Templates", "Manage templates and connect them to academic periods.", "Passed with Remark", "The tester reported an error when reusing the same template. Validation and locking rules were later verified; an identical Citra retest is not documented."),
        ("1.14", "Thesis Export", "Export thesis data to CSV.", "Passed with Remark", "The tester reported an export error. A dedicated export route and menu were added and technically verified."),
        ("1.15", "Final Documents and Library", "Open the final-document queue and confirm library availability.", "Passed with Remark", "The tester found missing final-document requirements and an empty queue. The current completion-to-library path is covered by automated tests."),
        ("1.16", "Sidang Schedule", "Set the thesis sidang date and time.", "Passed with Remark", "The tester saw a database error. The required database field and route rendering were later verified; production smoke testing remains recommended."),
        ("1.17", "Thesis Monitoring", "Open thesis monitoring from the Kaprodi workspace.", "Passed with Remark", "The tester saw a Blade syntax error. The view was corrected and route rendering now passes."),
        ("1.18", "Final Document Checklist", "Save a checklist as draft and later publish it.", "Passed with Remark", "The tester reported that the status stayed Draft. Current draft, publish, and locking rules are technically verified; an identical Citra retest is not documented."),
        ("1.19", "Final Document Monitoring", "Open final documents waiting for Kaprodi review.", "Passed with Remark", "The tester saw a date-format error. The page and date handling were corrected; an identical Citra retest is not documented."),
    ],
    "Dosen": [
        ("2.1", "Dashboard", "Open the Dosen dashboard and assigned thesis queue.", "Passed", "Verified"),
        ("2.2", "Assigned Theses", "View assigned theses and their details.", "Passed", "Verified"),
        ("2.3", "Guidance Records", "Add, edit, and delete guidance notes.", "Passed", "Verified"),
        ("2.4", "Sidang Request", "Submit, cancel, and submit a sidang request again.", "Passed with Remark", "The tester reported unclear feedback after submission. Laravel coverage and accepted TestSprite TC029 confirm cancellation and resubmission."),
        ("2.5", "Grading", "Enter grade items and publish the final grade.", "Passed", "Verified"),
        ("2.6", "Grade Unlock", "Request an unlock for a published grade.", "Passed", "Verified"),
        ("2.7", "Notifications", "Receive notifications for assignments, reviews, and requests.", "Passed", "Verified"),
    ],
    "Mahasiswa": [
        ("3.1", "Dashboard", "View the active thesis, document count, and guidance summary.", "Passed", "Verified"),
        ("3.2", "Thesis Management", "Create, edit, and delete an owned thesis record.", "Passed", "Verified"),
        ("3.3", "Documents", "Upload proposal and document versions, then delete owned files.", "Passed", "Verified"),
        ("3.4", "Guidance", "View guidance, respond to revisions, and export the logbook.", "Passed", "Verified"),
        ("3.5", "Final Submission", "Open and submit the final-document form only after grades are available.", "Passed with Remark", "The tester requested an eligibility guard. Current automated tests confirm that early access and early submission are rejected."),
        ("3.6", "Grades", "View the student's own published final grades.", "Passed", "Verified"),
        ("3.7", "Non-Thesis Records", "Manage an abstract, report, score, and publication link.", "Passed", "Verified"),
    ],
    "General": [
        ("4.1", "Authentication", "Use valid and invalid login details and confirm role-based access.", "Passed", "Verified"),
        ("4.2", "Notifications", "View realtime notifications and mark one or all as read.", "Passed", "Verified"),
        ("4.3", "Shared User Interface", "Check shared page headers, cards, tables, forms, and icons.", "Passed", "Verified"),
        ("4.4", "Public Library", "Open the public library list and thesis detail pages.", "Passed with Remark", "The tester reported that completed theses did not appear automatically. The corrected path is covered by five accepted browser tests and lifecycle tests."),
        ("4.5", "FAQ and Limitations", "Open the FAQ and Known Limitations pages.", "Passed", "Verified"),
    ],
}


def build_sequence_diagram():
    width, height = 1800, 1050
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_dir = Path("/System/Library/Fonts")
    regular_path = font_dir / "Supplemental/Arial.ttf"
    bold_path = font_dir / "Supplemental/Arial Bold.ttf"
    regular = ImageFont.truetype(str(regular_path), 29)
    small = ImageFont.truetype(str(regular_path), 25)
    bold = ImageFont.truetype(str(bold_path), 30)

    actors = [
        ("Mahasiswa", 150),
        ("Dosen", 500),
        ("Kaprodi", 850),
        ("TACLOUD", 1200),
        ("Public Library", 1570),
    ]
    navy = "#173B57"
    teal = "#147D78"
    gray = "#93A4AF"
    ink = "#26343D"

    for label, x in actors:
        box = (x - 125, 35, x + 125, 105)
        draw.rounded_rectangle(box, radius=12, fill=navy)
        text_box = draw.textbbox((0, 0), label, font=bold)
        draw.text((x - (text_box[2] - text_box[0]) / 2, 55), label, font=bold, fill="white")
        draw.line((x, 105, x, height - 35), fill=gray, width=3)

    messages = [
        (150, 1200, "Submit proposal"),
        (1200, 850, "Notify proposal review"),
        (850, 1200, "Approve and assign reviewers"),
        (1200, 500, "Notify assignment"),
        (500, 1200, "Record guidance"),
        (150, 1200, "Upload revisions"),
        (500, 1200, "Submit sidang request"),
        (1200, 850, "Notify sidang review"),
        (850, 1200, "Approve and set schedule"),
        (500, 1200, "Publish required grades"),
        (150, 1200, "Upload final document"),
        (850, 1200, "Complete thesis"),
        (1200, 1570, "Publish completed thesis"),
    ]
    y = 155
    for start, end, label in messages:
        color = teal if start != 1200 else navy
        draw.line((start, y, end, y), fill=color, width=4)
        direction = 1 if end > start else -1
        tip = end
        draw.polygon(
            [(tip, y), (tip - direction * 18, y - 10), (tip - direction * 18, y + 10)],
            fill=color,
        )
        text_box = draw.textbbox((0, 0), label, font=small)
        text_width = text_box[2] - text_box[0]
        text_x = (start + end) / 2 - text_width / 2
        draw.rectangle((text_x - 8, y - 33, text_x + text_width + 8, y - 2), fill="white")
        draw.text((text_x, y - 31), label, font=small, fill=ink)
        y += 67

    image.save(SEQUENCE_PNG, dpi=(180, 180))


def build_cover_visual():
    screenshot = Image.open(LOGIN_SCREENSHOT).convert("RGB")
    screenshot.thumbnail((1540, 870))
    canvas = Image.new("RGB", (1680, 1010), "#F5F8FA")
    shadow = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    x = (canvas.width - screenshot.width) // 2
    y = (canvas.height - screenshot.height) // 2
    shadow_draw.rounded_rectangle(
        (x - 24, y - 18, x + screenshot.width + 32, y + screenshot.height + 38),
        radius=34,
        fill=(23, 59, 87, 45),
    )
    canvas = Image.alpha_composite(canvas.convert("RGBA"), shadow)
    mask = Image.new("L", screenshot.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle((0, 0, screenshot.width, screenshot.height), radius=24, fill=255)
    canvas.paste(screenshot, (x, y), mask)
    border = ImageDraw.Draw(canvas)
    border.rounded_rectangle(
        (x - 2, y - 2, x + screenshot.width + 2, y + screenshot.height + 2),
        radius=26,
        outline="#D7E2E8",
        width=5,
    )
    canvas.convert("RGB").save(COVER_VISUAL, quality=95, dpi=(180, 180))


def setup(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (
        ("Title", 28, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 12.5, TEAL),
        ("Heading 3", 10.5, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "TACLOUD  |  Launch Readiness and Project Handover Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(7.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Version 3.1  |  27 July 2026  |  Page ")
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(footer, "PAGE")


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(UKWK_LOGO), width=Cm(2.0))
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TACLOUD")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Launch Readiness and\nProject Handover Report")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_para(doc, "Final project review for production launch", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(COVER_VISUAL), width=Cm(15.7))

    metadata = [
        ("Prepared for", "Head of Study Program (Kaprodi)"),
        ("Date", "27 July 2026"),
        ("Version", "3.1"),
        ("Recommendation", "READY FOR A CONTROLLED PRODUCTION LAUNCH"),
    ]
    t = doc.add_table(rows=len(metadata), cols=2)
    t.autofit = False
    for i, (label, value) in enumerate(metadata):
        shade(t.cell(i, 0), NAVY)
        shade(t.cell(i, 1), PALE)
        set_width(t.cell(i, 0), 4.2)
        set_width(t.cell(i, 1), 11.2)
        set_cell_text(t.cell(i, 0), label, bold=True, color=WHITE, size=8.5)
        set_cell_text(t.cell(i, 1), value, bold=i == 5, color=GREEN if i == 5 else INK, size=8.5)
    page_break(doc)


def executive_summary(doc):
    add_heading(doc, "Executive Summary", 1)
    callout(
        doc,
        "Launch recommendation",
        "TACLOUD is ready for a controlled production launch. No known issue currently prevents launch within the tested scope.",
        fill="EAF7F0",
        accent=GREEN,
    )
    add_para(
        doc,
        "TACLOUD supports the thesis process from proposal submission to supervision, sidang, grading, final-document submission, completion, and public library listing. The application separates access for Kaprodi, Dosen, and Mahasiswa.",
    )
    add_para(
        doc,
        "All 20 final test areas were completed. Seven accepted TestSprite browser tests passed. The current Laravel suite passed 135 tests with 457 assertions. Citra recorded all 38 manual QA checks as Passed.",
    )
    add_para(
        doc,
        "Eleven Citra checks included remarks. The related paths were later fixed or technically verified. However, the same manual test was not repeated by Citra for every remark. This is recorded as an evidence limitation.",
    )
    add_heading(doc, "Readiness at a glance", 2)
    table(
        doc,
        ["Evidence", "Result", "Status"],
        [
            ("Final test campaign", "20 of 20 scenarios completed", "Passed"),
            ("Accepted TestSprite tests", "7 of 7 browser scenarios passed", "Passed"),
            ("Laravel automated tests", "135 tests / 457 assertions", "Passed"),
            ("Focused workflow tests", "22 tests / 90 assertions", "Passed"),
            ("Critical release suites", "65 tests / 290 assertions", "Passed"),
            ("Authentication and import", "39 tests / 175 assertions", "Passed"),
            ("Dependency audits", "Composer: 0 advisories; npm: 0 vulnerabilities", "Passed"),
            ("Citra manual QA", "38 of 38 recorded Passed; 11 with remarks", "Passed with qualification"),
            ("Known launch blocker", "None within the verified scope", "Clear"),
        ],
        [5.2, 7.2, 3.3],
        font=8,
    )
    add_heading(doc, "Conditions before launch", 2)
    bullet(doc, "Complete the production deployment checklist.")
    bullet(doc, "Back up the database and private document storage, and confirm the rollback procedure.")
    bullet(doc, "Use HTTPS, secure cookies, and HSTS at the server or reverse proxy.")
    bullet(doc, "Run a short production smoke test for login, main role access, uploads, sidang, grades, completion, and library.")

    add_heading(doc, "Document Contents", 1)
    table(
        doc,
        ["Section", "Purpose"],
        [
            ("1. Project Overview", "Purpose, users, scope, and expected outputs"),
            ("2. Roles and Responsibilities", "Access and responsibility by user role"),
            ("3. How TACLOUD Works", "Main thesis flow and role handoffs"),
            ("4. Features Delivered", "Delivered functional areas"),
            ("5. Verification Method", "How launch readiness was checked"),
            ("6. Citra Manual QA", "All 38 checks with clear explanations"),
            ("7. Final Test Campaign", "All 20 final test scenarios"),
            ("8. Issues and Current Status", "Important findings and remediation"),
            ("9. Security and Production Readiness", "Security controls and release gates"),
            ("10. Launch Plan", "Deployment, rollback, and monitoring"),
        ],
        [5.7, 10.0],
        font=8,
    )
def project_and_flow(doc):
    add_heading(doc, "1. Project Overview", 1)
    add_para(doc, "TACLOUD is a web application for managing the academic thesis process in one system. It reduces separate files, unclear status updates, and manual handoffs between Mahasiswa, Dosen, and Kaprodi.")
    table(
        doc,
        ["Item", "Description"],
        [
            ("Main purpose", "Manage the thesis process from proposal to final publication."),
            ("Primary users", "Kaprodi, Dosen, and Mahasiswa."),
            ("Public user", "Can search and view information about completed theses."),
            ("Main outputs", "Current thesis status, documents, guidance records, sidang requests, schedules, grades, completion records, exports, and library entries."),
            ("Authentication", "Password login and institutional Google login for @widyakarya.ac.id accounts."),
        ],
        [4.0, 11.7],
        font=8.2,
    )

    add_heading(doc, "2. Roles and Responsibilities", 1)
    table(
        doc,
        ["Activity", "Kaprodi", "Dosen", "Mahasiswa"],
        [
            ("Master data and periods", "Manage", "View when needed", "View own context"),
            ("Proposal", "Review and decide", "View assigned thesis", "Create and upload"),
            ("Supervisor and examiner", "Assign", "Perform assigned role", "View assignment"),
            ("Guidance", "Monitor", "Record guidance", "View and upload revision"),
            ("Sidang request", "Approve or reject", "Submit or cancel pending request", "View progress"),
            ("Grades", "View recap and approve unlock", "Enter, publish, and request unlock", "View own grades"),
            ("Final document", "Review and complete thesis", "View assigned thesis", "Submit when eligible"),
            ("Library", "Monitor", "Public access", "Public access"),
        ],
        [4.2, 3.7, 4.1, 3.7],
        font=7.5,
    )

    page_break(doc)
    add_heading(doc, "3. How TACLOUD Works", 1)
    table(
        doc,
        ["Step", "Responsible Role", "Main Action and Result"],
        [
            ("1", "Mahasiswa", "Creates a thesis record and uploads the proposal."),
            ("2", "Kaprodi", "Reviews the proposal and approves or returns it for revision."),
            ("3", "Kaprodi", "Assigns supervisors and examiners."),
            ("4", "Dosen and Mahasiswa", "Record guidance, comments, revisions, and document versions."),
            ("5", "Primary Supervisor", "Submits the sidang request when requirements are met."),
            ("6", "Kaprodi", "Approves or rejects the request and sets the schedule."),
            ("7", "Assigned Dosen", "Enter and publish the required grades."),
            ("8", "Mahasiswa", "Uploads the final document after all required grades are available."),
            ("9", "Kaprodi", "Checks the final document and published grades, then completes the thesis."),
            ("10", "System", "Shows the completed thesis in the public library."),
        ],
        [1.5, 4.4, 9.8],
        font=8,
    )
    callout(
        doc,
        "Important workflow correction",
        "The old Dosen final-document approval step was removed because it was not part of the required business flow. Completion now uses the latest final document and published grades from all assigned reviewers.",
        fill=LIGHT,
        accent=TEAL,
    )
    add_heading(doc, "3.1 Main Role Handoffs", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SEQUENCE_PNG), width=Cm(15.7))
    add_para(
        doc,
        "The diagram shows who starts each action and how TACLOUD passes the work to the next responsible role.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=GRAY,
        after=2,
    )
    page_break(doc)


def features_and_method(doc):
    add_heading(doc, "4. Features Delivered", 1)
    table(
        doc,
        ["Functional Area", "Delivered Capability", "Status"],
        [
            ("Authentication", "Password login, logout, session protection, and institutional Google login.", "Delivered"),
            ("Master Data", "Dosen, Mahasiswa, academic years, periods, grading formats, and CSV imports.", "Delivered"),
            ("Thesis Management", "Owned thesis records, proposal review, assignments, phases, and monitoring.", "Delivered"),
            ("Guidance", "Guidance notes, revisions, document versions, and logbook exports.", "Delivered"),
            ("Sidang", "Requests, cancellation while pending, Kaprodi decision, and scheduling.", "Delivered"),
            ("Grading", "Draft grades, publication and lock, recap, and controlled unlock requests.", "Delivered"),
            ("Final Documents", "Eligibility checks, final upload, Kaprodi review, normal and historical completion.", "Delivered"),
            ("Document Templates", "Required-document lists, period links, draft, publish, duplicate, and locking rules.", "Delivered"),
            ("Notifications", "Personal notifications, read status, internal links, and realtime delivery support.", "Delivered"),
            ("Reports and Export", "Thesis CSV, grade recap, monitoring, and guidance logbook export.", "Delivered"),
            ("Public Library", "Search and detail pages for completed theses.", "Delivered"),
            ("Non-Thesis Records", "Student management of abstracts, reports, scores, and publication links.", "Delivered"),
        ],
        [4.1, 9.6, 2.0],
        font=7.7,
    )

    add_heading(doc, "5. Verification Method", 1)
    add_para(doc, "Launch readiness was checked with several methods. Each method was used for the type of risk it can verify best.")
    table(
        doc,
        ["Method", "What It Verified"],
        [
            ("Citra manual QA", "Real user tasks for Kaprodi, Dosen, Mahasiswa, and shared functions."),
            ("TestSprite browser tests", "Accepted user flows in a real browser, including library, sidang cancellation, and the two-advisor sidang flow."),
            ("Laravel/Pest tests", "Server rules, permissions, ownership, validation, files, security, workflows, and notifications."),
            ("Build and cache checks", "Frontend build, route cache, configuration cache, view cache, and Laravel optimization."),
            ("Dependency audit", "Known security advisories in Composer and npm packages."),
            ("Code and graph review", "Current structure, affected files, workflow connections, and removed obsolete behavior."),
        ],
        [5.0, 10.7],
        font=8,
    )
    page_break(doc)


def qa_section(doc):
    add_heading(doc, "6. Manual QA by Citra", 1)
    add_para(doc, "Citra recorded all 38 checks as Passed. Eleven checks included remarks. The tables keep the original result separate from later technical verification.")
    table(
        doc,
        ["Test Group", "Total", "Passed", "Passed with Remark"],
        [
            ("Kaprodi", "19", "11", "8"),
            ("Dosen", "7", "6", "1"),
            ("Mahasiswa", "7", "6", "1"),
            ("General", "5", "4", "1"),
            ("Total", "38", "27", "11"),
        ],
        [5.0, 3.0, 3.7, 4.0],
        font=8.2,
    )
    callout(
        doc,
        "How to read the status",
        "Passed is Citra's original result. Passed with Remark means Citra recorded Passed but also described a problem or improvement. Current Status explains later evidence without pretending that Citra repeated every test.",
    )

    for index, (group, rows) in enumerate(QA.items(), start=1):
        if group == "Mahasiswa":
            page_break(doc)
        add_heading(doc, f"6.{index} {group} QA", 2)
        table(
            doc,
            ["No.", "Test Scenario", "What Was Tested", "Tester Result", "Current Status"],
            rows,
            [1.0, 3.3, 4.8, 2.3, 4.3],
            font=6.8,
        )
    page_break(doc)


def final_campaign(doc):
    add_heading(doc, "7. Final Test Campaign", 1)
    add_para(doc, "The final campaign covered 20 test areas. Passed means that the final verification found no launch-blocking issue in that area.")
    table(
        doc,
        ["No.", "Test Scenario", "What Was Tested", "Final Result"],
        TEST_SCENARIOS,
        [1.0, 4.2, 8.5, 2.0],
        font=6.6,
    )
    add_para(
        doc,
        "Some scenarios found earlier defects or risks. The issues were corrected and tested again. The most important examples are summarized in Section 8; detailed evidence remains in the technical evidence package.",
        color=GRAY,
    )

    add_heading(doc, "Accepted TestSprite Browser Results", 2)
    table(
        doc,
        ["Browser Test Group", "Coverage", "Result", "Evidence Note"],
        [
            ("Public Library", "Five browse, detail, and navigation scenarios", "5/5 Passed", "Accepted browser assertions"),
            ("Dosen Sidang Cancellation", "Cancel a pending request and submit it again", "1/1 Passed", "Laravel tests also verify the success message"),
            ("Two-Advisor Sidang Flow", "Both supervisors submit; Kaprodi approves; phase advances", "1/1 Passed", "Database and Laravel workflow evidence support the result"),
            ("Old TC023", "Removed Dosen final-approval workflow", "Excluded", "No longer represents the corrected business flow"),
        ],
        [3.5, 5.3, 2.2, 4.7],
        font=7.4,
    )


def issues_and_security(doc):
    add_heading(doc, "8. Important Issues and Current Status", 1)
    table(
        doc,
        ["Issue Found", "Action Taken", "Verification", "Current Status"],
        [
            ("The old final-document approval design could not reach normal completion.", "Removed the unused approval model, UI status, tests, and database table. Completion now checks final documents and published grades.", "Focused workflow and complete lifecycle tests passed.", "Closed"),
            ("Dosen sidang requests remained visible after submission.", "Added clear pending-state handling, cancellation, and resubmission.", "Accepted TestSprite TC029 and Laravel feature tests passed.", "Closed"),
            ("Unsafe shared-table HTML could allow XSS.", "Escaped normal cell content and limited raw HTML to trusted action content.", "Focused XSS tests and frontend build passed.", "Closed"),
            ("Unsafe sort direction could cause a server error.", "Limited sort direction to accepted values.", "Query-security tests passed.", "Closed"),
            ("Proposal upload accepted an incorrect document phase.", "Restricted the endpoint to the proposal phase.", "Upload and phase-validation tests passed.", "Closed"),
            ("A duplicate route name blocked production optimization.", "Kept one main route name and renamed or removed duplicates.", "Route cache and optimization passed.", "Closed"),
            ("Some Citra remarks do not have an identical manual retest.", "Kept the original remarks and linked later technical verification.", "Evidence is clear but not identical for every manual step.", "Non-blocking evidence limitation"),
            ("Developer test accounts and a predictable import password were present in production paths.", "Removed the Akun Test shortcut and all runtime known credentials. CSV import now requires an explicit password of at least eight characters.", "Fresh-install, authentication, import, full-suite, and reference-audit checks passed.", "Closed"),
        ],
        [4.1, 5.1, 4.2, 2.3],
        font=7.2,
    )

    add_heading(doc, "9. Security and Production Readiness", 1)
    table(
        doc,
        ["Control Area", "Verified Result", "Production Requirement"],
        [
            ("Authentication", "Password and institutional Google login protections passed; no test-account shortcut or known production credential remains.", "Use production Google credentials and HTTPS callback URLs."),
            ("Role and ownership", "Cross-role access and record ownership tests passed.", "Keep exact role middleware and current authorization checks."),
            ("Forms and input", "CSRF, validation, mass-assignment, query, and sorting checks passed.", "Do not disable middleware or validation in production."),
            ("Files", "Type, path, phase, ownership, preview, and download checks passed.", "Back up private storage and verify permissions."),
            ("Browser protection", "Anti-sniff, frame, and referrer headers passed.", "Enable HSTS at the HTTPS server or reverse proxy."),
            ("Dependencies", "Independent live-registry audits on 27 July 2026 found 0 Composer advisories and 0 npm vulnerabilities.", "Repeat audits before each release."),
            ("Build and cache", "Vite 8.1.5 built 58 modules; route cache, configuration cache, views, and optimize passed.", "Build from the approved release commit."),
            ("Session", "Session controls passed in test scope.", "Set secure cookies in production."),
        ],
        [4.0, 6.3, 5.4],
        font=7.5,
    )


def launch_and_signoff(doc):
    add_heading(doc, "10. Launch Plan", 1)
    callout(
        doc,
        "Final recommendation",
        "Proceed with a controlled production launch after the checklist below is completed. This recommendation does not mean zero risk.",
        fill="EAF7F0",
        accent=GREEN,
    )
    add_heading(doc, "10.1 Deployment Checklist", 2)
    for item in [
        "Freeze the approved release commit and record the application version.",
        "Back up the database and private document storage.",
        "Confirm that the backup can be restored.",
        "Set the production URL, database, mail, Google login, queue, and notification settings.",
        "Use HTTPS and secure session cookies.",
        "Enable HSTS at the server or reverse proxy after HTTPS is confirmed.",
        "Install production dependencies and build the approved frontend assets.",
        "Run database migrations and confirm their result.",
        "Seed production master data only; the verified seeder creates no users or known credentials.",
        "Create the first Kaprodi through institutional Google login, then promote only that exact account with a scoped operator command.",
        "Build Laravel configuration, route, event, and view caches.",
        "Start and monitor queue, scheduler, and realtime services where used.",
        "Run smoke tests for login, all three roles, uploads, sidang, grades, final completion, export, notifications, and library.",
        "Record the launch time, responsible person, and first monitoring result.",
    ]:
        bullet(doc, "☐ " + item)

    add_heading(doc, "10.2 Backup and Rollback", 2)
    table(
        doc,
        ["Stage", "Required Action"],
        [
            ("Before deployment", "Create a consistent database and document-storage backup. Keep the previous application release."),
            ("Rollback trigger", "Repeated server errors, failed login for a main role, migration failure, broken uploads, or data mismatch."),
            ("Application rollback", "Return to the previous approved release, rebuild caches, and restart background workers."),
            ("Data rollback", "Restore the consistent backup when a database migration or data change cannot be safely reversed."),
            ("After rollback", "Check login, role access, record counts, documents, grades, notifications, and library entries."),
        ],
        [4.0, 11.7],
        font=8,
    )

    page_break(doc)
    add_heading(doc, "10.3 Post-Launch Monitoring", 2)
    table(
        doc,
        ["Time", "What to Monitor", "Owner"],
        [
            ("First 2 hours", "Login, server errors, migrations, queue, uploads, downloads, and response time.", "Developer / Operator"),
            ("First day", "Proposal, sidang, grading, completion, notifications, and data consistency.", "Developer / Kaprodi"),
            ("First week", "Error trends, failed jobs, storage growth, user feedback, and unusual access failures.", "Developer"),
            ("Second week", "Review incidents and decide priorities for remaining non-blocking improvements.", "Developer / Kaprodi"),
        ],
        [3.0, 9.6, 3.1],
        font=8,
    )

    page_break(doc)
    add_heading(doc, "Appendix A - Evidence Sources", 1)
    for source in [
        "TACLOUD current source code, routes, controllers, models, migrations, views, and feature tests.",
        "Code-review graph incrementally refreshed on 27 July 2026 with no errors.",
        "TestSprite_Final_Campaign_Evidence.md.",
        "TACLOUD_Current_Product_Flows_and_Roles.md.",
        "Citra public manual QA Google Sheet: 38 checklist rows.",
        "Local Laravel, Vite, Composer, npm, route-cache, and optimization results.",
        "Accepted TestSprite dashboard results and generated browser-test artifacts.",
    ]:
        bullet(doc, source)

    add_heading(doc, "Appendix B - Evidence Limits", 1)
    bullet(doc, "Eleven Citra checks contain remarks. Identical manual retesting by Citra is not documented for every remark.")
    bullet(doc, "The old TC023 result is excluded because the related Dosen final-approval workflow was removed.")
    bullet(doc, "The TestSprite success-message assertion for Dosen cancellation was not present; Laravel feature tests verify the session feedback.")
    bullet(doc, "HSTS and secure cookies depend on the final production HTTPS configuration.")
    bullet(doc, "The Vite warning for /images/login-wave.jpg is non-blocking because the file is available from the public directory at runtime.")
    bullet(doc, "Dependency audit results are time-sensitive. The report uses independent live-registry results from 27 July 2026 and audits must be repeated before release.")
    bullet(doc, "Fresh-install proof used a disposable SQLite database: 57 migrations, zero users, four roles, one department, and two study programs. Shared MySQL was not reset.")
    bullet(doc, "Never run migrate:fresh on shared hosting or production.")


def main():
    assert sum(len(rows) for rows in QA.values()) == 38
    assert len(TEST_SCENARIOS) == 20
    build_sequence_diagram()
    build_cover_visual()
    doc = Document()
    setup(doc)
    document_settings(doc)
    cover(doc)
    executive_summary(doc)
    project_and_flow(doc)
    features_and_method(doc)
    qa_section(doc)
    final_campaign(doc)
    issues_and_security(doc)
    launch_and_signoff(doc)
    doc.core_properties.title = "TACLOUD Launch Readiness and Project Handover Report"
    doc.core_properties.subject = "Simple professional project handover and launch recommendation"
    doc.core_properties.author = "TACLOUD Project Team"
    doc.core_properties.keywords = "TACLOUD, launch readiness, project handover, QA, TestSprite"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
