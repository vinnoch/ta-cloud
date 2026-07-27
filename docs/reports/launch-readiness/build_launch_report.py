from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = Path(__file__).resolve().parent
DOCX = OUT / "Laporan_Kesiapan_Peluncuran_TACLOUD.docx"
QA_CSV = Path("/private/tmp/citra_qa.csv")

NAVY = "173B57"
TEAL = "147D78"
GOLD = "D8A23A"
LIGHT = "EAF2F5"
PALE = "F5F8FA"
GREEN = "217A4A"
AMBER = "9A6700"
RED = "A33A3A"
GRAY = "66737D"
WHITE = "FFFFFF"
INK = "26343D"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, *, bold=False, color=INK, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text or "-"))
    r.bold = bold
    r.font.name = "Aptos"
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")


def table(doc, headers, rows, widths, *, font=7.6):
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    t.style = "Table Grid"
    set_repeat_header(t.rows[0])
    keep_row_together(t.rows[0])
    for i, (header, width) in enumerate(zip(headers, widths)):
        set_width(t.rows[0].cells[i], width)
        shade(t.rows[0].cells[i], NAVY)
        set_cell_text(t.rows[0].cells[i], header, bold=True, color=WHITE, size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, row_data in enumerate(rows):
        cells = t.add_row().cells
        keep_row_together(t.rows[-1])
        for i, (value, width) in enumerate(zip(row_data, widths)):
            set_width(cells[i], width)
            if row_idx % 2:
                shade(cells[i], PALE)
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font, align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, *, bold_lead=None, color=None, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    if color:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def callout(doc, label, text, fill=LIGHT, accent=TEAL):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    c = t.cell(0, 0)
    shade(c, fill)
    margins(c, 150, 180, 150, 180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def page_break(doc):
    doc.add_page_break()


def status_color(status):
    s = status.lower()
    if "terblokir" in s or "blocking" in s or "gap" in s:
        return RED
    if "lulus" in s or "tertutup" in s or "siap" in s:
        return GREEN
    if "terbatas" in s or "non-blocking" in s or "perlu" in s:
        return AMBER
    return RED


def read_qa():
    with QA_CSV.open(encoding="utf-8-sig", newline="") as fh:
        rows = list(csv.reader(fh))
    groups = {"Kaprodi": [], "Dosen": [], "Mahasiswa": [], "Bersama": []}
    group = None
    for row in rows[1:]:
        row += [""] * (7 - len(row))
        no, area, task, status, notes, _, feedback = [x.strip() for x in row[:7]]
        marker = f"{no} {area}"
        if marker.startswith("1.") and "Checklist" in marker:
            group = "Kaprodi"
        elif marker.startswith("2.") and "Checklist" in marker:
            group = "Dosen"
        elif marker.startswith("3.") and "Checklist" in marker:
            group = "Mahasiswa"
        elif marker.startswith("4.") and "Checklist" in marker:
            group = "Bersama"
        elif group and no and task:
            groups[group].append({"no": no, "area": area, "task": task, "status": status,
                                  "notes": notes, "feedback": feedback})
    return groups


RESOLUTIONS = {
    "1.11": "Ditindaklanjuti pada alur keputusan sidang; route/controller dan notifikasi terverifikasi oleh suite workflow. Tidak ada retest manual Citra yang tercatat.",
    "1.13": "Asosiasi template-periode kini dijaga oleh validasi/aturan penguncian. Catatan konflik tetap dicatat; retest manual identik belum terdokumentasi.",
    "1.14": "Backend ekspor CSV tersedia melalui menu/route khusus Kaprodi. Disposisi: ditutup pada kode; perlu smoke test produksi setelah deploy.",
    "1.15": "Alur template, final submission, final review, dan publikasi library kini terhubung. Feature test membuktikan skripsi selesai tampil di library.",
    "1.16": "Kolom jadwal sidang ditambahkan melalui migrasi dan route rendering/optimize lulus. Disposisi: ditutup teknis; retest manual produksi dianjurkan.",
    "1.17": "Syntax Blade diperbaiki pada kampanye poin 2; seluruh route-render test kini lulus.",
    "1.18": "Kode mendukung draft/publish dan penguncian template. Belum ada bukti retest manual spesifik untuk skenario perubahan draft ke publish; non-blocking, wajib smoke test.",
    "1.19": "Casting/tampilan tanggal dan route render kini lulus. Belum ada retest manual Citra spesifik; perlu smoke test halaman Dokumen Final.",
    "2.4": "UX diperjelas dengan pembatalan pengajuan pending. TC029 membatalkan lalu mengajukan ulang dan lulus; success message diverifikasi oleh feature test, bukan assertion browser.",
    "3.5": "Akses final submission sekarang dikunci sampai nilai lengkap. Lima feature test mencakup eligible, prematur, ownership, dan penyimpanan final.",
    "4.4": "Skripsi selesai otomatis masuk library; diuji pada workflow/cross-role dan XSS library.",
}


POINTS = [
    ("1", "Isolasi lingkungan dan data uji", "Menyiapkan lingkungan produksi-semu terpisah.",
     "Konfigurasi, build, health check, seed MySQL terisolasi.",
     "SQLite bersih gagal karena schema squash; MySQL terisolasi dipakai. Google onboarding kemudian diubah agar akun domain sah dapat dibuat saat login pertama.",
     "Perbaikan schema test dan Google login diuji ulang.",
     "Seed 34 pengguna, 55 migrasi, debug nonaktif, HTTP redirect login normal.", "LULUS"),
    ("2", "Baseline statis, suite, dan build", "Mendapat baseline regresi dan build.",
     "Composer/Pest penuh, Vite build, route rendering.",
     "Awal 19 lulus/74 gagal; schema test, test stale, ikon/Blade, dan assertion UI bermasalah.",
     "Schema/test/Blade dipulihkan oleh task terkait; rerun penuh.",
     "93 test/298 assertion saat penutupan poin; baseline hijau.", "LULUS"),
    ("3", "Infrastruktur TestSprite", "Membuktikan browser remote dan sinkronisasi hasil.",
     "Generate plan; run terpotong: TC019 lalu TC018/022/030/040.",
     "Tunnel Yamux membuat progres lokal 0/N meski remote lanjut; run penuh dihindari untuk hemat kredit.",
     "Menunggu worker natural exit dan menjalankan batch kecil.",
     "5/5 browser test library lulus; artefak lokal tersinkron.", "LULUS"),
    ("4", "Autentikasi", "Menguji password, Google institusi, sesi, throttle, dan disclosure.",
     "Auth + GoogleSecurity suite; tambah coverage invalid password/logout/rate limit/disabled user/enumeration.",
     "Coverage eksplisit awal belum lengkap.",
     "Lima test fokus ditambahkan tanpa perubahan perilaku inti.",
     "32 test/153 assertion; domain @widyakarya.ac.id dan sesi aman.", "LULUS"),
    ("5", "Otorisasi berbasis peran", "Mencegah akses silang Kaprodi/Dosen/Mahasiswa.",
     "Role middleware, cross-role, route, forged write, stale-role session.",
     "Probe awal memakai ID palsu sehingga 404 mendahului middleware; fixture diperbaiki.",
     "Test memakai record nyata dan menambah forged-write/stale-session.",
     "36 test/95 assertion; tidak ada bypass peran.", "LULUS"),
    ("6", "Ownership dan IDOR", "Mencegah pengguna satu peran mengubah record milik peer.",
     "Direct request pada bimbingan, dokumen, submission, grading.",
     "Dosen assigned dapat update/hapus bimbingan dosen lain.",
     "Controller mewajibkan reviewer_id = pengguna terautentikasi.",
     "35 test/95 assertion; cross-owner update/delete 403.", "LULUS"),
    ("7", "CSRF", "Membuktikan request tanpa token ditolak.",
     "POST/PUT/DELETE/login dengan middleware CSRF aktif plus kontrol token valid.",
     "Tidak ada defect; perlu membedakan 419 sengaja dari form nyata.",
     "Tidak ada patch.",
     "Request palsu 419; token valid diproses.", "LULUS"),
    ("8", "Validasi dan mass assignment", "Mencegah field privileged dan nilai batas.",
     "Malformed/boundary/forged fields lintas peran.",
     "Coverage ownership/status/role/grade boundary belum eksplisit.",
     "Tiga regression test ditambahkan.",
     "71 test/275 assertion; eskalasi field ditolak.", "LULUS"),
    ("9", "Keamanan upload/download", "Menguji MIME, ukuran, phase/path, storage, ownership.",
     "Direct upload forgery dan unauthorized download.",
     "Mahasiswa dapat mengirim phase skripsi_final lewat endpoint proposal.",
     "Phase dibatasi exact proposal; path tetap disanitasi.",
     "26 test/77 assertion; forged phase ditolak.", "LULUS"),
    ("10", "XSS dan escaping", "Mencegah stored/reflected/DOM XSS.",
     "Payload pada library, shared table, dan search suggestions.",
     "Shared table merender semua cell sebagai raw HTML.",
     "Cell biasa di-escape; HtmlString hanya untuk action; DOM memakai textContent.",
     "4 test/17 assertion dan build lulus.", "LULUS"),
    ("11", "Query dan sort manipulation", "Menguji SQL payload, search, sort, filter.",
     "QuerySecurity + route checks.",
     "Sort direction berbahaya menyebabkan 500, bukan SQL execution.",
     "Direction dinormalisasi asc/desc pada dua jalur.",
     "18 test/62 assertion; payload menjadi data.", "LULUS"),
    ("12", "Header dan sesi browser", "Memastikan header anti-sniff/frame/referrer dan cookie posture.",
     "Respons normal, redirect, error; audit config produksi.",
     "Header baseline belum ada.",
     "Middleware global menambah nosniff, SAMEORIGIN, strict-origin-when-cross-origin.",
     "18 test/83 assertion. Secure cookie/HSTS wajib di HTTPS; CSP ditunda.", "LULUS TERBATAS"),
    ("13", "Workflow Mahasiswa", "Memvalidasi proposal hingga final/non-skripsi.",
     "Seluruh suite Mahasiswa + guard final prematur.",
     "Guard transisi final belum diuji langsung.",
     "Regression memastikan file/fase tak berubah sebelum nilai lengkap.",
     "32 test/88 assertion.", "LULUS"),
    ("14", "Workflow Dosen", "Memvalidasi assignment, bimbingan, sidang, nilai, unlock.",
     "Seluruh suite Dosen dan test permohonan sidang.",
     "Coverage submit sidang kritis belum ada.",
     "Test primary advisor vs reviewer tidak berwenang ditambahkan.",
     "17 test/49 assertion.", "LULUS"),
    ("15", "Workflow Kaprodi dan legacy", "Memvalidasi keputusan dan penyelesaian normal/historis.",
     "Kaprodi suite + direct-action transition probes.",
     "Endpoint dapat menyelesaikan skripsi tanpa bukti workflow.",
     "Normal butuh dokumen final, assignment, dan nilai published; jalur legacy Kaprodi khusus periode nonaktif.",
     "Business-flow correction menghapus approval dokumen final yang tidak terpakai. Focused workflow 22 test/90 assertion lulus.", "LULUS"),
    ("16", "Lifecycle lintas peran", "Membuktikan handoff sampai library.",
     "Integration + workflow transition + library assertion.",
     "Coverage completion-to-library ditambahkan.",
     "Assertion publikasi library ditambah.",
     "Cross-role completion dan publikasi library lulus; accepted browser coverage total 7/7.", "LULUS"),
    ("17", "Notifikasi", "Menguji penerima, unread, read/all, link, ownership.",
     "Notification feature dan cross-role producers.",
     "Tidak ada defect.",
     "Tidak ada patch.",
     "7 test/22 assertion.", "LULUS"),
    ("18", "Aksesibilitas dan UI", "Menguji dialog, keyboard, build, rendered routes.",
     "Markup audit, view cache, build; browser runtime gagal init.",
     "27 modal tanpa semantics; focus trap/Escape/restoration tidak konsisten.",
     "Semantics + satu helper JS native untuk keyboard/focus.",
     "4 fokus/15 assertion; full 126/418 saat penutupan; browser-level gap dicatat.", "LULUS TERBATAS"),
    ("19", "Kesiapan produksi dan cache", "Memastikan route/config/view cache dan build dapat dibuat.",
     "optimize, route cache, audits, route uniqueness.",
     "Nama route add-periode duplikat memblokir optimize.",
     "Nama canonical dipertahankan; legacy diberi nama unik; duplicate route lain dihapus.",
     "9 test/37 assertion; optimize dan route cache lulus.", "LULUS"),
    ("20", "Gerbang rilis final", "Menjalankan seluruh gate dan audit dependensi.",
     "Full suite, build, optimize, audit npm/Composer, validate, diff.",
     "Audit awal menemukan 20 advisory pada 11 package.",
     "44 update patch/minor terkontrol; Socialite dipin ^5.27.",
     "Saat kampanye: 127/421 dan zero advisory. Verifikasi kini: 133/452, build/optimize/audit hijau.", "LULUS"),
]


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.1)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(1.8)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 10),
        ("Heading 1", 18, NAVY, 18, 8),
        ("Heading 2", 13, TEAL, 13, 6),
        ("Heading 3", 11, NAVY, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        h = sec.header
        p = h.paragraphs[0]
        p.text = "TACLOUD  |  Laporan Kesiapan Peluncuran v2.1"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
        f = sec.footer
        p = f.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Dokumen v2.1  |  20 Juli 2026  |  Halaman ")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(GRAY)
        add_field(p, "PAGE")


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(75)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LAPORAN KESIAPAN\nPELUNCURAN")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TACLOUD")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    add_para(doc, "Sistem Manajemen Tugas Akhir", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=26)
    t = doc.add_table(rows=5, cols=2)
    t.autofit = False
    metadata = [
        ("Tujuan", "Penilaian kesiapan go-live berbasis bukti"),
        ("Ditujukan kepada", "Ketua Program Studi (Kaprodi)"),
        ("Tanggal", "20 Juli 2026"),
        ("Versi", "2.1"),
        ("Status", "SIAP DALAM CAKUPAN TERUJI"),
    ]
    for i, (label, value) in enumerate(metadata):
        shade(t.cell(i, 0), NAVY)
        shade(t.cell(i, 1), PALE)
        set_cell_text(t.cell(i, 0), label, bold=True, color=WHITE, size=9)
        set_cell_text(t.cell(i, 1), value, bold=(i == 4), color=GREEN if i == 4 else INK, size=9)
        set_width(t.cell(i, 0), 4)
        set_width(t.cell(i, 1), 11.5)
    add_para(doc, "\nDisusun dari kode terkini, graph arsitektur, project tracking, QA manual Citra, "
             "kampanye final 20 poin, dan verifikasi ulang release gate.", align=WD_ALIGN_PARAGRAPH.CENTER,
             color=GRAY, after=0)
    page_break(doc)


def executive(doc):
    add_heading(doc, "Ringkasan Eksekutif", 1)
    callout(doc, "Kesimpulan", "TACLOUD siap diluncurkan dalam cakupan yang telah diuji. Business-flow correction "
            "menghapus workflow approval dokumen final yang tidak dipakai; completion kini memeriksa dokumen final "
            "dan nilai reviewer published.", fill="EAF7F0", accent=GREEN)
    add_para(doc, "Bukti terkini menunjukkan 133 test backend lulus dengan 452 assertion, build produksi Vite "
             "berhasil, cache konfigurasi/rute/view berhasil dibuat, dan Composer audit tidak menemukan advisory. "
             "Graph kode dibangun ulang pada 20 Juli 2026: 361 file, 905 node, 1.601 edge setelah post-processing "
             "(1.608 pada hasil build awal), tanpa error.")
    add_para(doc, "QA manual Citra mencatat 38/38 baris berstatus Passed. Namun beberapa baris berisi defect atau "
             "gap; laporan ini mempertahankan catatan asli dan memisahkannya dari bukti perbaikan kemudian. "
             "Sebagian skenario telah ditutup oleh automated tests atau TestSprite, sedangkan beberapa tetap "
             "memerlukan smoke test produksi.")
    add_heading(doc, "Bukti utama", 2)
    bullet(doc, "Kampanye final: 20/20 skenario ditutup.")
    bullet(doc, "TestSprite diterima: 7/7 browser case — 5 library, 1 cancellation/resubmit, 1 alur sidang dua pembimbing.")
    bullet(doc, "Focused workflow: 22 test/90 assertion; full suite: 133/452.")
    bullet(doc, "QA manual Citra: 38/38 recorded Passed; 11 baris berkomentar dipertahankan sebagai PASS*.")
    add_heading(doc, "Keterbatasan tambahan", 2)
    bullet(doc, "Mahasiswa dapat memilih periode yang ada tetapi server belum mewajibkan periode aktif.")
    bullet(doc, "Library publik menampilkan metadata, tetapi route dokumen final masih mensyaratkan autentikasi.")
    bullet(doc, "CSP belum diterapkan karena inline script/style masih luas; header nosniff, SAMEORIGIN, dan referrer policy sudah aktif.")
    bullet(doc, "HSTS harus diatur pada server/reverse proxy HTTPS; SESSION_SECURE_COOKIE=true wajib di produksi.")
    bullet(doc, "Vite memberi warning /images/login-wave.jpg; file ada di public/images dan memang diselesaikan saat runtime.")
    bullet(doc, "Aksesibilitas modal teruji pada kode/markup/build; browser runtime audit penuh sempat tidak tersedia.")
    bullet(doc, "Beberapa defect QA Citra tidak memiliki retest manual identik; ditandai untuk smoke test pra-go-live.")
    add_heading(doc, "Daftar Isi", 1)
    table(doc, ["Bagian", "Halaman"], [
        ("Ringkasan Eksekutif", "2"),
        ("Bab 1 — Persyaratan Produk dan Alur Operasional", "3"),
        ("Bab 2 — Hasil QA Manual oleh Citra", "6"),
        ("Bab 3 — Kampanye Final 20 Poin", "10"),
        ("Kesimpulan dan Rekomendasi Peluncuran", "17"),
        ("Lampiran — Sumber Bukti", "18"),
    ], [13.8, 1.9], font=8.5)
    page_break(doc)


def chapter1(doc):
    add_heading(doc, "Bab 1 — Persyaratan Produk dan Alur Operasional", 1)
    add_heading(doc, "1.1 Tujuan, ruang lingkup, dan keluaran", 2)
    add_para(doc, "TACLOUD mendigitalisasi lifecycle tugas akhir: data pengguna dan periode, proposal, reviewer, "
             "bimbingan, permohonan sidang, penilaian, dokumen final, penyelesaian, notifikasi, ekspor, dan library.")
    table(doc, ["Area", "Cakupan aktif", "Keluaran utama"], [
        ("Akademik", "Tahun akademik, periode, Dosen, Mahasiswa, import CSV", "Data operasional terstruktur"),
        ("Skripsi", "Proposal, assignment, bimbingan, sidang, nilai, final review", "Status dan bukti lifecycle"),
        ("Dokumen", "Versi, preview/download terotorisasi, template/checklist final", "Arsip privat dan dokumen final"),
        ("Pelaporan", "Rekap nilai, export skripsi, logbook CSV/PDF", "Bahan monitoring dan administrasi"),
        ("Publik", "Library skripsi selesai", "Daftar/detail karya akhir"),
        ("Pendukung", "Notifikasi database/realtime, pencarian, FAQ", "Umpan balik lintas peran"),
    ], [3.0, 8.0, 4.7], font=8)

    add_heading(doc, "1.2 Peran, izin, dan tanggung jawab", 2)
    table(doc, ["Aktivitas", "Kaprodi", "Dosen", "Mahasiswa"], [
        ("Data master/periode/template", "Kelola global", "Lihat sesuai kebutuhan", "Lihat konteks sendiri"),
        ("Proposal", "Review/approve/reject", "Lihat assigned", "Buat dan unggah"),
        ("Reviewer/jadwal", "Tetapkan/atur", "Jalankan assignment", "Lihat"),
        ("Bimbingan", "Monitor", "Buat/ubah/hapus milik sendiri", "Lihat/respons/revisi"),
        ("Permohonan sidang", "Approve/reject", "Pembimbing 1 submit/cancel pending", "Lihat progres"),
        ("Penilaian", "Format, rekap, unlock", "Isi/publish-lock/request unlock", "Lihat nilai sendiri"),
        ("Dokumen final", "Validasi/selesaikan; legacy khusus", "Lihat dokumen assigned", "Kirim jika eligible"),
        ("Library", "Monitor", "Akses publik", "Akses publik"),
    ], [5.2, 3.4, 3.8, 3.3], font=7.5)

    add_heading(doc, "1.3 Alur end-to-end", 2)
    flow = [
        ("1", "Mahasiswa", "Buat skripsi dan unggah proposal"),
        ("2", "Kaprodi", "Review proposal; setujui/tolak; atur sidang proposal"),
        ("3", "Kaprodi", "Tetapkan pembimbing/penguji"),
        ("4", "Dosen ↔ Mahasiswa", "Bimbingan, catatan, revisi, versi dokumen"),
        ("5", "Pembimbing 1", "Ajukan sidang; dapat batalkan selama masih submitted"),
        ("6", "Kaprodi", "Setujui/tolak permohonan; atur jadwal sidang skripsi"),
        ("7", "Dosen", "Isi dan publish-lock nilai; request unlock bila perlu"),
        ("8", "Mahasiswa", "Unggah dokumen final setelah nilai lengkap"),
        ("9", "Kaprodi", "Validasi dokumen dan nilai reviewer; nyatakan selesai"),
        ("10", "Sistem", "Tampilkan skripsi selesai di library"),
    ]
    table(doc, ["Tahap", "Penanggung jawab", "Handoff/hasil"], flow, [1.5, 4.2, 10.0], font=8)
    callout(doc, "Koreksi alur", "Approval dokumen final Dosen dihapus sebagai workflow yang tidak digunakan. "
            "Kaprodi menyelesaikan skripsi setelah dokumen final tersedia dan semua nilai reviewer sudah published.",
            fill="EAF7F0", accent=GREEN)
    callout(doc, "Jalur historis", "Untuk lulusan periode nonaktif/ditutup, Kaprodi dapat mengunggah dokumen final "
            "tervalidasi dan menyelesaikan record lewat jalur legacy terpisah. Jalur ini ditolak pada periode aktif.")

    add_heading(doc, "1.3.1 Flowchart keputusan penyelesaian", 3)
    table(doc, ["Urutan", "Keputusan", "Ya", "Tidak"], [
        ("A", "Dokumen final terbaru tersedia?", "Periksa periode", "Tetap review"),
        ("B", "Periode aktif?", "Periksa nilai semua reviewer", "Kaprodi dapat pilih jalur legacy"),
        ("C", "Semua nilai sidang published?", "Kaprodi nyatakan selesai", "Tetap review"),
        ("D", "Skripsi selesai?", "Masuk katalog library", "Tidak dipublikasikan"),
    ], [1.4, 6.4, 4.1, 3.8], font=7.8)

    add_heading(doc, "1.4 Fase dan transisi valid", 2)
    table(doc, ["Fase", "Masuk ketika", "Keluar ketika"], [
        ("proposal", "Record dibuat / proposal ditolak untuk revisi", "Proposal disetujui"),
        ("sidang_proposal", "Proposal siap sidang", "Hasil/assignment mengarahkan ke bimbingan"),
        ("bimbingan_skripsi", "Pembimbing ditetapkan", "Permohonan sidang siap diproses"),
        ("sidang_skripsi", "Seluruh syarat permohonan disetujui", "Sidang selesai; revisi bila perlu"),
        ("revisi_sidang_skripsi", "Ada perbaikan pascasidang", "Nilai dan dokumen final lengkap"),
        ("review_dokumen_final", "Final submission diterima", "Kaprodi memvalidasi semua bukti"),
        ("skripsi_selesai", "Normal completion atau legacy nonaktif", "Status terminal; masuk library"),
    ], [4.0, 6.0, 5.7], font=7.8)

    add_heading(doc, "1.5 Aturan bisnis dan keamanan", 2)
    rules = [
        ("Autentikasi", "Google hanya domain persis @widyakarya.ac.id. Login pertama dapat membuat akun least-privilege; sesi diregenerasi. Password login tetap tersedia."),
        ("Otorisasi", "Role middleware memisah workspace. Dosen dibatasi assignment; Mahasiswa dibatasi ownership; record nested diverifikasi."),
        ("Dokumen", "Private storage dan controller access; MIME/ukuran divalidasi; endpoint proposal hanya menerima phase proposal; final memakai endpoint khusus."),
        ("Nilai", "Rentang 0–100; publish-lock mencegah perubahan biasa; Dosen request unlock, Kaprodi memproses."),
        ("Periode", "Satu konteks aktif; data berelasi tidak dihapus sembarang. Legacy completion hanya untuk periode nonaktif/ditutup."),
        ("Penyelesaian", "Normal completion butuh final document, reviewer assignments, dan nilai sidang published untuk setiap assignment."),
        ("Notifikasi", "Penerima dideduplikasi; link dinormalisasi internal; read state hanya milik pengguna."),
        ("Library", "Metadata publik hanya untuk skripsi selesai dan user content di-escape. Preview/download final masih berada di route berautentikasi."),
    ]
    table(doc, ["Aturan", "Implementasi operasional"], rules, [3.2, 12.5], font=8)

    add_heading(doc, "1.6 Alur pendukung", 2)
    bullet(doc, "Master data: CRUD/import CSV Dosen dan Mahasiswa, tahun akademik, periode, format penilaian, template dokumen.")
    bullet(doc, "Bimbingan: pertemuan, catatan dosen/mahasiswa, revisi, dokumen yang direview, ekspor logbook CSV/PDF.")
    bullet(doc, "Notifikasi: database + realtime, unread count, read satu/semua, tautan internal.")
    bullet(doc, "Ekspor: rekap skripsi Kaprodi dan logbook Mahasiswa.")
    bullet(doc, "Non-skripsi: CRUD ringkasan, abstrak, laporan, nilai, dan tautan publikasi milik Mahasiswa; monitoring Kaprodi.")
    bullet(doc, "Library: index/detail publik untuk karya selesai; dokumen tetap melalui kontrol akses yang sesuai.")
    page_break(doc)


def chapter2(doc, groups):
    add_heading(doc, "Bab 2 — Hasil QA Manual oleh Citra", 1)
    add_para(doc, "Sumber: Google Sheet publik Citra, dibaca 18 Juli 2026. Total 38 checklist: 19 Kaprodi, "
             "7 Dosen, 7 Mahasiswa, 5 bersama/auth/UI. Kolom Hasil asli dipertahankan; catatan defect tidak "
             "diubah menjadi pass bersih. Kolom Resolusi membedakan bukti kemudian.")
    table(doc, ["Kelompok", "Jumlah", "Status asli"], [
        ("Kaprodi", "19", "19 Passed; 9 baris memuat catatan/gap"),
        ("Dosen", "7", "7 Passed; 1 baris memuat catatan UX"),
        ("Mahasiswa", "7", "7 Passed; 1 baris memuat gap validasi"),
        ("Bersama", "5", "5 Passed; 1 baris memuat gap publikasi"),
        ("Total", "38", "38 Passed; catatan tetap direkonsiliasi"),
    ], [5.0, 2.5, 8.2], font=8)

    for group_name, items in groups.items():
        add_heading(doc, f"2.{list(groups).index(group_name)+1} {group_name}", 2)
        rows = []
        for item in items:
            note = item["notes"] or "Tidak ada catatan."
            if item["feedback"]:
                note += " | Feedback developer asli: " + item["feedback"]
            resolution = RESOLUTIONS.get(item["no"], "Tidak ada defect tercatat; diterima sebagai hasil QA manual asli.")
            disposition = "Ditutup/terverifikasi" if item["no"] in {"1.14", "1.15", "1.17", "2.4", "3.5", "4.4"} else (
                "Perlu smoke test" if item["no"] in {"1.11", "1.13", "1.16", "1.18", "1.19"} else "Passed manual"
            )
            rows.append((item["no"], item["area"] + " — " + item["task"], item["status"],
                         note, resolution + " | Disposisi: " + disposition))
        table(doc, ["No.", "Tugas", "Hasil asli", "Catatan tester", "Respons/resolusi & posisi kini"],
              rows, [1.2, 4.7, 1.7, 4.1, 4.0], font=6.7)

    add_heading(doc, "2.5 Isu penting dan status", 2)
    table(doc, ["Tema", "Temuan awal", "Bukti penanganan", "Status"], [
        ("Sidang", "Approve error dan UX request tidak berubah", "Workflow suite; cancellation TC029; feature session feedback", "Ditutup, smoke produksi"),
        ("Template/final", "Konflik template, draft/publish, antrean final", "Controller rules, final-submission/workflow tests", "Sebagian perlu smoke manual"),
        ("Jadwal", "QueryException saat jadwal", "Migrasi kolom + route/optimize hijau", "Ditutup teknis"),
        ("Monitoring", "Syntax Blade dan format tanggal", "Point 2 Blade fix; route rendering hijau", "Ditutup teknis"),
        ("Ekspor", "CSV error/backend gap", "Route/controller export tersedia", "Ditutup teknis"),
        ("Library", "Final tidak otomatis muncul", "Point 16 completion-to-library test", "Ditutup otomatis"),
    ], [3.0, 4.5, 5.3, 2.9], font=7.4)
    page_break(doc)


def chapter3(doc):
    add_heading(doc, "Bab 3 — Kampanye Final 20 Poin", 1)
    add_para(doc, "Kampanye memadukan TestSprite browser dengan Laravel/Pest, code review, build, route/cache, "
             "audit keamanan, dan inspeksi aksesibilitas. Karena TestSprite berbasis browser tidak cocok untuk "
             "semua kontrol server, metode dibedakan secara eksplisit.")
    add_heading(doc, "3.1 Pemisahan metode", 2)
    table(doc, ["Metode", "Pemakaian"], [
        ("TestSprite browser", "7/7 accepted: library 5/5, cancellation/resubmit 1/1, sidang dua pembimbing 1/1."),
        ("Laravel/Pest", "Auth, role, IDOR, CSRF, validasi, upload, XSS response, query, headers, workflow, notifications."),
        ("Static/code review", "Root-cause inspection, routes/controllers/Blade/JS, graph impact."),
        ("Build/cache/audit", "Vite, optimize/route cache, Composer/npm audit, manifest validation."),
        ("Accessibility", "Markup regression, shared keyboard helper, Blade cache/build; browser runtime limitation dicatat."),
    ], [4.0, 11.7], font=8)

    for point in POINTS:
        no, title, objective, method, initial, fix, rerun, final = point
        add_heading(doc, f"3.{no} Poin {no} — {title}", 2)
        t = doc.add_table(rows=6, cols=2)
        t.autofit = False
        entries = [
            ("Tujuan/ruang lingkup", objective),
            ("Metode", method),
            ("Hasil awal/risiko", initial),
            ("Remediasi", fix),
            ("Hasil rerun", rerun),
            ("Status final", final),
        ]
        for i, (label, value) in enumerate(entries):
            shade(t.cell(i, 0), NAVY if i < 5 else status_color(final))
            shade(t.cell(i, 1), PALE if i % 2 == 0 else WHITE)
            set_width(t.cell(i, 0), 4.2)
            set_width(t.cell(i, 1), 11.5)
            set_cell_text(t.cell(i, 0), label, bold=True, color=WHITE, size=8)
            set_cell_text(t.cell(i, 1), value, bold=(i == 5), size=8)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_heading(doc, "3.21 Retest khusus pembatalan permohonan sidang Dosen", 2)
    add_para(doc, "TestSprite TC029 semula memberi false positive karena hanya memeriksa URL. Pemeriksaan data "
             "menemukan fixture drift: request seeded berstatus approved sehingga tombol memang disembunyikan. "
             "Fixture dikoreksi menjadi submitted secara deterministik, tanpa melemahkan guard UI/controller.")
    table(doc, ["Aspek", "Bukti"], [
        ("Browser", "TC029 mengklik Batalkan pengajuan sidang, melihat tombol Ajukan kembali, mengirim ulang, kembali ke /dosen/skripsi/3, dan memastikan label pembatalan muncul lagi."),
        ("Hasil", "Passed — TestSprite run f5d15013-838c-4426-92d8-3eb050491f49."),
        ("Gap browser", "Assertion success message tidak dibuat oleh TestSprite."),
        ("Pelengkap", "Laravel feature test memverifikasi deletion, redirect/session success, visibility, processed-state guard, dan lecturer lain 404."),
    ], [3.6, 12.1], font=8)
    add_para(doc, "URL hasil: https://www.testsprite.com/dashboard/mcp/tests/d01df0fb-6922-40ee-97be-68c33cc39f0d/f5d15013-838c-4426-92d8-3eb050491f49",
             color=TEAL)

    add_heading(doc, "3.22 Koreksi business flow dokumen final", 2)
    add_para(doc, "Verifikasi workspace aktif menunjukkan workflow approval dokumen final lama tidak diperlukan "
             "oleh aturan bisnis. Model, test, status UI, dan tabel approval dihapus; migration telah dijalankan. "
             "TC023 lama dikeluarkan karena tidak lagi merepresentasikan workflow produk.")
    table(doc, ["Aspek", "Bukti terkini"], [
        ("Final submission", "Menyimpan versi dokumen final dan memindahkan fase ke review_dokumen_final."),
        ("Kaprodi completion", "Mensyaratkan dokumen final terbaru dan nilai sidang published untuk setiap assignment."),
        ("Migration", "Tabel final_document_approvals yang obsolete dihapus; migration tercatat Ran."),
        ("Focused workflow", "22 test/90 assertion lulus."),
        ("TC023", "Dikeluarkan dari accepted browser count karena workflow lama telah dihapus."),
        ("Kesimpulan", "Workflow teruji dapat mencapai skripsi_selesai dan publikasi library."),
    ], [4.0, 11.7], font=8)

    add_heading(doc, "3.23 Bukti agregat terkini", 2)
    table(doc, ["Gate", "Hasil 18 Juli 2026", "Status"], [
        ("Code-review graph", "20 Juli: 361 file; 905 node; 1.601 edge post-process; 0 error", "Lulus"),
        ("Laravel/Pest", "133 test; 452 assertion", "Lulus"),
        ("Focused workflow", "22 test; 90 assertion", "Lulus"),
        ("TestSprite", "7/7 accepted browser cases", "Lulus"),
        ("Frontend", "Vite 8.0.8; 56 modul; build selesai", "Lulus"),
        ("Laravel optimize", "Config, events, routes, views cached", "Lulus"),
        ("Composer audit", "No security vulnerability advisories found", "Lulus"),
        ("Workflow normal", "Final submission → Kaprodi completion → library", "Lulus"),
        ("Warning", "/images/login-wave.jpg diselesaikan runtime; file public tersedia", "Non-blocking"),
    ], [4.1, 8.3, 3.3], font=8)
    page_break(doc)


def conclusion(doc):
    add_heading(doc, "Kesimpulan dan Rekomendasi Peluncuran", 1)
    callout(doc, "Rekomendasi", "LANJUTKAN GO-LIVE DALAM CAKUPAN TERUJI, dengan checklist deployment, backup, "
            "HTTPS, secure cookie, HSTS, migration check, dan smoke test produksi. Tidak ada blocker terbuka pada "
            "workflow yang telah diverifikasi.", fill="EAF7F0", accent=GREEN)
    add_heading(doc, "Klasifikasi risiko", 2)
    table(doc, ["Kelas", "Kondisi"], [
        ("Blocking", "Tidak ada blocker terbuka pada workflow yang diuji."),
        ("Launch-impacting", "Pemilihan periode belum mewajibkan active; scope file library publik belum diputuskan; override fase Kaprodi terlalu luas untuk alur normal."),
        ("Non-blocking", "CSP belum aktif; HSTS/secure cookie bergantung deploy; warning asset runtime; beberapa QA manual butuh smoke retest; browser accessibility audit penuh terbatas."),
        ("Tidak diterima", "Peluncuran tanpa backup/rollback, HTTPS, secure cookie, HSTS, migration check, atau smoke test role utama."),
    ], [3.5, 12.2], font=8)

    add_heading(doc, "Checklist wajib sebelum deployment", 2)
    checklist = [
        "Pastikan migration penghapusan final_document_approvals berstatus Ran pada produksi.",
        "Jalankan E2E smoke: Mahasiswa submit final → Kaprodi complete → library.",
        "Wajibkan periode aktif/eligible pada create skripsi atau dokumentasikan aturan pengecualian terotorisasi.",
        "Putuskan scope library: metadata-only atau public final-file delivery dengan route khusus aman.",
        "Bekukan commit/artifact rilis dan catat checksum/versi.",
        "Backup database dan storage dokumen; uji restore pada lingkungan terpisah.",
        "Set APP_ENV=production, APP_DEBUG=false, APP_URL=https://..., key/secret produksi aman.",
        "Set SESSION_SECURE_COOKIE=true; HTTPS wajib end-to-end.",
        "Aktifkan HSTS pada web server/reverse proxy setelah HTTPS tervalidasi.",
        "Jalankan composer install --no-dev --classmap-authoritative dan npm build/artifact yang disetujui.",
        "Jalankan migration --force setelah backup; catat waktu dan hasil.",
        "Jalankan optimize/route cache; pastikan queue/realtime/scheduler aktif.",
        "Pastikan permission private storage/log/cache benar; public symlink sesuai kebutuhan.",
        "Smoke test Google @widyakarya.ac.id, tiap role, proposal, sidang, nilai, final, export, library, notifikasi.",
        "Smoke test ulang QA Citra prioritas: sidang approve, template draft/publish, jadwal, Dokumen Final.",
        "Pantau error log, queue failure, response 5xx, auth failure, storage, dan notifikasi.",
    ]
    for item in checklist:
        bullet(doc, "☐ " + item)

    add_heading(doc, "Backup dan rollback", 2)
    table(doc, ["Tahap", "Tindakan"], [
        ("Sebelum deploy", "Snapshot database + storage; simpan artifact versi sebelumnya; dokumentasikan migration."),
        ("Trigger rollback", "5xx berulang, migrasi gagal, auth lintas role gagal, upload/download rusak, data integrity anomali."),
        ("Rollback aplikasi", "Alihkan ke artifact sebelumnya; clear/rebuild cache; restart worker."),
        ("Rollback data", "Gunakan migration down hanya bila aman dan teruji; jika tidak, restore snapshot konsisten database+storage."),
        ("Verifikasi", "Login tiga role, record count, dokumen, nilai, notifikasi, library, dan audit log."),
    ], [3.8, 11.9], font=8)

    add_heading(doc, "Rencana monitoring pascapeluncuran", 2)
    table(doc, ["Waktu", "Pantauan", "Pemilik"], [
        ("0–2 jam", "5xx, auth/Google callback, migration, queue, upload/download, response time", "Developer + operator"),
        ("Hari 1", "Workflow proposal/sidang/nilai/final, notification delivery, data mismatch", "Developer + Kaprodi"),
        ("Hari 2–7", "Error trend, queue retry, storage growth, user feedback, access denied anomaly", "Developer"),
        ("Minggu 2", "Review incident/non-blocking backlog; putuskan CSP nonce/hash rollout", "Developer + Kaprodi"),
    ], [3.0, 9.3, 3.4], font=8)

    add_heading(doc, "Pernyataan akhir", 2)
    add_para(doc, "Laporan ini tidak menyatakan “zero risk”. Rekomendasi siap dalam cakupan teruji didasarkan pada "
             "graph/kode terkini, 20/20 kampanye, 7/7 accepted browser cases, 133 test/452 assertion, focused workflow "
             "22/90, dan 38/38 QA manual recorded Passed. Sebelas baris Citra tetap ditandai PASS* karena memuat "
             "remark; retest manual identik tidak terdokumentasi untuk semuanya.")

    add_heading(doc, "Lampiran — Sumber bukti", 1)
    sources = [
        "Code-review graph TACLOUD, full rebuild 20 Juli 2026.",
        "Evidence package: docs/reports/launch-readiness/evidence/TACLOUD_Current_Product_Flows_and_Roles.md.",
        "Evidence package: docs/reports/launch-readiness/evidence/TestSprite_Final_Campaign_Evidence.md.",
        "Rekonsiliasi langsung workspace aktif terhadap graph, routes, controllers, views, models, dan tests.",
        "TestSprite task 019f682d-88e8-7932-bfd3-c1939880b4f3 dan artefak testsprite_tests/.",
        "Citra Manual QA Google Sheet: docs.google.com/spreadsheets/d/1ufbAdoYl6v7-Boa2d29jBe5GZ5NVv42B999I-Asnp90/.",
        "PRD.md, README.md, plan.md, routes, controllers, models, migrations, Blade/JS, dan tests/Feature.",
        "Verifikasi lokal 20 Juli 2026: composer test, npm run build, php artisan optimize, Composer/npm audit, validate, diff check.",
    ]
    for src in sources:
        bullet(doc, src)


def document_settings(doc):
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)


def main():
    groups = read_qa()
    assert {k: len(v) for k, v in groups.items()} == {"Kaprodi": 19, "Dosen": 7, "Mahasiswa": 7, "Bersama": 5}
    doc = Document()
    setup_styles(doc)
    document_settings(doc)
    cover(doc)
    executive(doc)
    chapter1(doc)
    chapter2(doc, groups)
    chapter3(doc)
    conclusion(doc)
    doc.core_properties.title = "Laporan Kesiapan Peluncuran TACLOUD"
    doc.core_properties.subject = "Launch readiness untuk Kaprodi"
    doc.core_properties.author = "Tim TACLOUD"
    doc.core_properties.keywords = "TACLOUD, launch readiness, QA, TestSprite, Kaprodi"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
